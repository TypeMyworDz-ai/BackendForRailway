import logging
import sys
import asyncio
import subprocess
import os
import json
import base64
from contextlib import asynccontextmanager
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Response, Request, Form
from fastapi.middleware.cors import CORSMiddleware
import tempfile
import uuid
from datetime import datetime, timedelta
import requests
from pydub import AudioSegment
from pydantic import BaseModel
from typing import Optional, List
import httpx
from docx import Document
from docx.shared import Inches
from io import BytesIO
from fastapi.responses import StreamingResponse
import re
import anthropic
import openai

import google.generativeai as genai

import firebase_admin
from firebase_admin import credentials, firestore, initialize_app
from google.cloud.firestore_v1.base_query import FieldFilter


logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

logger.info("=== STARTING FASTAPI APPLICATION (MAIN BACKEND) ===")

# Service Names
TYPEMYWORDZ1_NAME = "TypeMyworDz1" # AssemblyAI
TYPEMYWORDZ2_NAME = "TypeMyworDz2" # OpenAI Whisper
DEEPGRAM_NAME = "Deepgram" # Deepgram
TYPEMYWORDZ_AI_NAME = "TypeMyworDz AI" # Anthropic Claude / OpenAI GPT / Google Gemini

# Admin email addresses
ADMIN_EMAILS = ['typemywordz@gmail.com', 'gracenyaitara@gmail.com']
# Dedicated AssemblyAI Tester
ASSEMBLYAI_TESTER_EMAIL = 'njokigituku@gmail.com'

def install_ffmpeg():
    try:
        subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
        logger.info("ffmpeg is already installed")
    except (subprocess.CalledProcessError, FileNotFoundError):
        logger.info("Installing ffmpeg... (This might not be strictly necessary if pydub uses a pre-installed one on Railway)")
        try:
            subprocess.run(['apt-get', 'update'], check=True)
            subprocess.run(['apt-get', 'install', '-y', 'ffmpeg'], check=True)
            logger.info("ffmpeg installed successfully")
        except subprocess.CalledProcessError as e:
            logger.error(f"Failed to install ffmpeg: {e}")

install_ffmpeg()
logger.info("Loading environment variables...")

ASSEMBLYAI_API_KEY = os.environ.get("ASSEMBLYAI_API_KEY")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
PAYSTACK_SECRET_KEY = os.environ.get("PAYSTACK_SECRET_KEY")
PAYSTACK_PUBLIC_KEY = os.environ.get("PAYSTACK_PUBLIC_KEY")
PAYSTACK_WEBHOOK_SECRET = os.environ.get("PAYSTACK_WEBHOOK_SECRET")
OPENAI_WHISPER_SERVICE_RAILWAY_URL = os.environ.get("OPENAI_WHISPER_SERVICE_RAILWAY_URL")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
FIREBASE_ADMIN_SDK_CONFIG_BASE64 = os.environ.get("FIREBASE_ADMIN_SDK_CONFIG_BASE64")
DEEPGRAM_SERVICE_RAILWAY_URL = os.environ.get("DEEPGRAM_SERVICE_RAILWAY_URL")

logger.info(f"DEBUG: --- Environment Variable Check (main.py) ---")
logger.info(f"DEBUG: ASSEMBLYAI_API_KEY loaded value: {bool(ASSEMBLYAI_API_KEY)}")
logger.info(f"DEBUG: ANTHROPIC_API_KEY loaded value: {bool(ANTHROPIC_API_KEY)}")
logger.info(f"DEBUG: OPENAI_API_KEY (for GPT if direct) loaded value: {bool(OPENAI_API_KEY)}")
logger.info(f"DEBUG: PAYSTACK_SECRET_KEY loaded value: {bool(PAYSTACK_SECRET_KEY)}")
logger.info(f"DEBUG: PAYSTACK_PUBLIC_KEY loaded value: {bool(PAYSTACK_PUBLIC_KEY)}")
logger.info(f"DEBUG: PAYSTACK_WEBHOOK_SECRET loaded value: {bool(PAYSTACK_WEBHOOK_SECRET)}")
logger.info(f"DEBUG: OPENAI_WHISPER_SERVICE_RAILWAY_URL loaded value: {bool(OPENAI_WHISPER_SERVICE_RAILWAY_URL)}")
logger.info(f"DEBUG: GEMINI_API_KEY loaded value: {bool(GEMINI_API_KEY)}")
logger.info(f"DEBUG: FIREBASE_ADMIN_SDK_CONFIG_BASE64 loaded value: {bool(FIREBASE_ADMIN_SDK_CONFIG_BASE64)}")
logger.info(f"DEBUG: DEEPGRAM_SERVICE_RAILWAY_URL loaded value: {bool(DEEPGRAM_SERVICE_RAILWAY_URL)}")
logger.info(f"DEBUG: Admin emails configured: {ADMIN_EMAILS}")
logger.info(f"DEBUG: AssemblyAI Tester email: {ASSEMBLYAI_TESTER_EMAIL}")
logger.info(f"DEBUG: --- End Environment Variable Check (main.py) ---")

if not ASSEMBLYAI_API_KEY:
    logger.error(f"{TYPEMYWORDZ1_NAME} API Key environment variable not set! {TYPEMYWORDZ1_NAME} will not work as primary or fallback.")

if not ANTHROPIC_API_KEY:
    logger.warning(f"{TYPEMYWORDZ_AI_NAME} (Anthropic) API Key environment variable not set! Anthropic AI features will be disabled.")

if not OPENAI_API_KEY:
    logger.warning(f"OPENAI_API_KEY (for GPT if direct) environment variable not set! Direct OpenAI GPT calls disabled.")

if not OPENAI_WHISPER_SERVICE_RAILWAY_URL:
    logger.error(f"{TYPEMYWORDZ2_NAME} (OpenAI Whisper & GPT) Service URL not configured! OpenAI transcription and GPT formatting will be disabled.")

if not GEMINI_API_KEY:
    logger.warning("Google Gemini API Key environment variable not set! Google Gemini AI features will be disabled.")

if not DEEPGRAM_SERVICE_RAILWAY_URL:
    logger.error(f"{DEEPGRAM_NAME} Service URL not configured! {DEEPGRAM_NAME} will not work as primary or fallback.")

if not FIREBASE_ADMIN_SDK_CONFIG_BASE64:
    logger.error("FIREBASE_ADMIN_SDK_CONFIG_BASE64 environment variable not set! Firebase Admin SDK features (user/revenue updates) will be disabled.")
else:
    try:
        decoded_json_test = base64.b64decode(FIREBASE_ADMIN_SDK_CONFIG_BASE64).decode('utf-8')
        parsed_json_test = json.loads(decoded_json_test)
        logger.info(f"DIAGNOSTIC (Runtime): Firebase config decoded and parsed successfully. Project ID: {parsed_json_test.get('project_id', 'N/A')}, Client Email: {parsed_json_test.get('client_email', 'N/A')}")
    except Exception as e:
        logger.error(f"DIAGNOSTIC (Runtime): ERROR: Failed to decode/parse FIREBASE_ADMIN_SDK_CONFIG_BASE64 at runtime: {e}")


if not PAYSTACK_SECRET_KEY:
    logger.warning("PAYSTACK_SECRET_KEY environment variable not set! Paystack features will be disabled.")

if PAYSTACK_SECRET_KEY:
    logger.info("Paystack configuration found - payment verification enabled")
else:
    logger.warning("Paystack configuration missing - payment verification disabled")

logger.info("Environment variables loaded successfully")

# NEW: Initialize Firebase Admin SDK
db = None
if FIREBASE_ADMIN_SDK_CONFIG_BASE64:
    try:
        service_account_info = json.loads(base64.b64decode(FIREBASE_ADMIN_SDK_CONFIG_BASE64).decode('utf-8'))
        cred = credentials.Certificate(service_account_info)
        initialize_app(cred)
        db = firestore.client()
        logger.info("Firebase Admin SDK initialized successfully.")
    except Exception as e:
        logger.error(f"Error initializing Firebase Admin SDK: {e}")
else:
    logger.warning("Firebase Admin SDK config is missing, Firestore operations will not be available.")

claude_client = None
if ANTHROPIC_API_KEY:
    try:
        claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        logger.info(f"{TYPEMYWORDZ_AI_NAME} (Anthropic) client initialized successfully.")
    except Exception as e:
        logger.error(f"Error initializing {TYPEMYWORDZ_AI_NAME} (Anthropic) client: {e}")
else:
    logger.warning(f"{TYPEMYWORDZ_AI_NAME} (Anthropic) API key is missing, Claude client will not be initialized.")

# Google Gemini Client initialization
gemini_client = None
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        gemini_client = genai.GenerativeModel('models/gemini-pro-latest')
        logger.info(f"Google Gemini client initialized successfully.")
    except Exception as e:
        logger.error(f"Error initializing Google Gemini client: {e}")
else:
    logger.warning(f"Google Gemini API key is missing, client will not be initialized.")


def is_paid_ai_user(user_plan: str) -> bool:
    paid_plans_for_ai = ['Three-Day Plan', 'One-Week Plan', 'Monthly Plan', 'Yearly Plan']
    return user_plan in paid_plans_for_ai

def is_admin_user(user_email: str) -> bool:
    """Check if user is an admin based on email address"""
    if not user_email:
        return False
    return user_email.lower().strip() in [email.lower() for email in ADMIN_EMAILS]

def get_transcription_services(user_plan: str, speaker_labels_enabled: bool, user_email: str = None):
    """
    Logic for service selection based on new rules, including Deepgram.
    - OpenAI: First option for weekly subscribers, yearly, and Admins (Admins gets this logic no matter what plans they have subscribed to). Fallback is Assembly > Deepgram.
    - Assembly: First option for free users. Fallback Deepgram only (free users don't get TypeMyworDz Assistant) All instances of speaker tags requests: First option Deepgram, fallback Assembly.
    - Deepgram: First option for three-day and monthly plans users. Fallback is OpenAI > Assembly. All instances of speaker tags requests: First option Assembly, fallback Deepgram.
    - njokigituku@gmail.com will now be using AssemblyAI, which means for them there is no fallback if Assembly fails.
    """
    
    is_admin = is_admin_user(user_email) if user_email else False
    is_assemblyai_tester = (user_email and user_email.lower().strip() == ASSEMBLYAI_TESTER_EMAIL.lower())

    # --- Initialize tiers ---
    tier_1 = None
    tier_2 = None
    tier_3 = None
    reason = "default_logic"

    # --- Dedicated AssemblyAI Tester Logic ---
    if is_assemblyai_tester:
        tier_1 = "assemblyai" 
        reason = "dedicated_assemblyai_tester"
        # No fallbacks for the dedicated tester as per requirement
        return {
            "tier_1": tier_1,
            "tier_2": None,
            "tier_3": None,
            "reason": reason
        }

    # --- Speaker Labels Logic (Global Override) ---
    # As per request: "All instances of speaker tags requests: First option Assembly, fallback Deepgram."
    if speaker_labels_enabled:
        tier_1 = "assemblyai"
        tier_2 = "deepgram"
        tier_3 = None  # No OpenAI for speaker tags as per request for this flow
        reason = "speaker_labels_requested_prioritizing_assemblyai"
    
    # --- Plan-based Logic (if speaker labels are not enabled or already set) ---
    # OpenAI: First option for weekly subscribers, yearly, and Admins. Fallback is Assembly > Deepgram.
    elif is_admin or user_plan in ['One-Week Plan', 'Yearly Plan']:
        tier_1 = "openai_whisper"
        tier_2 = "assemblyai"
        tier_3 = "deepgram"
        reason = f"admin_or_{user_plan.lower().replace(' ', '_')}_prioritizing_openai"
    # Assembly: First option for free users. Fallback Deepgram only (free users don't get TypeMyworDz Assistant)
    elif user_plan == 'free': # FIX: Changed 'Free' to 'free'
        tier_1 = "assemblyai"
        tier_2 = "deepgram"
        tier_3 = None # No TypeMyworDz Assistant (OpenAI fallback) for free users
        reason = f"{user_plan.lower().replace(' ', '_')}_prioritizing_assemblyai"
    # Deepgram: First option for three-day and monthly plans users. Fallback is OpenAI > Assembly.
    elif user_plan in ['Three-Day Plan', 'Monthly Plan']:
        tier_1 = "deepgram"
        tier_2 = "openai_whisper"
        tier_3 = "assemblyai"
        reason = f"{user_plan.lower().replace(' ', '_')}_prioritizing_deepgram"
    
    # --- Dynamic adjustment based on service availability ---
    final_tiers_list = []
    
    # Tier 1
    if tier_1 == "assemblyai" and ASSEMBLYAI_API_KEY:
        final_tiers_list.append("assemblyai")
    elif tier_1 == "openai_whisper" and OPENAI_WHISPER_SERVICE_RAILWAY_URL:
        final_tiers_list.append("openai_whisper")
    elif tier_1 == "deepgram" and DEEPGRAM_SERVICE_RAILWAY_URL:
        final_tiers_list.append("deepgram")

    # Tier 2 (only if not already in Tier 1 and is available)
    if tier_2 == "assemblyai" and ASSEMBLYAI_API_KEY and "assemblyai" not in final_tiers_list:
        final_tiers_list.append("assemblyai")
    elif tier_2 == "openai_whisper" and OPENAI_WHISPER_SERVICE_RAILWAY_URL and "openai_whisper" not in final_tiers_list:
        final_tiers_list.append("openai_whisper")
    elif tier_2 == "deepgram" and DEEPGRAM_SERVICE_RAILWAY_URL and "deepgram" not in final_tiers_list:
        final_tiers_list.append("deepgram")

    # Tier 3 (only if not already in Tier 1 or 2 and is available)
    if tier_3 == "assemblyai" and ASSEMBLYAI_API_KEY and "assemblyai" not in final_tiers_list:
        final_tiers_list.append("assemblyai")
    elif tier_3 == "openai_whisper" and OPENAI_WHISPER_SERVICE_RAILWAY_URL and "openai_whisper" not in final_tiers_list:
        final_tiers_list.append("openai_whisper")
    elif tier_3 == "deepgram" and DEEPGRAM_SERVICE_RAILWAY_URL and "deepgram" not in final_tiers_list:
        final_tiers_list.append("deepgram")

    # Ensure the list does not exceed 3 tiers and fills Nones if less than 3
    return {
        "tier_1": final_tiers_list[0] if len(final_tiers_list) > 0 else None,
        "tier_2": final_tiers_list[1] if len(final_tiers_list) > 1 else None,
        "tier_3": final_tiers_list[2] if len(final_tiers_list) > 2 else None,
        "reason": reason
    }


class PaystackVerificationRequest(BaseModel):
    reference: str

class PaystackInitializationRequest(BaseModel):
    email: str
    amount: float
    plan_name: str
    user_id: str
    country_code: str
    callback_url: str
    update_admin_revenue: Optional[bool] = False

class PaystackWebhookRequest(BaseModel):
    event: str
    data: dict

class CreditUpdateRequest(BaseModel):
    email: str
    plan_name: str
    amount: float
    currency: str
    duration_hours: Optional[int] = None
    duration_days: Optional[int] = None

class FormattedWordDownloadRequest(BaseModel):
    transcription_html: str
    filename: Optional[str] = "transcription.docx"

class UserAIRequest_Pydantic(BaseModel):
    transcript: str
    user_prompt: str
    model: str = "claude-3-haiku-20240307"
    max_tokens: int = 1000

class AdminAIFormatRequest_Pydantic(BaseModel):
    transcript: str
    formatting_instructions: str = "Format the transcript for readability, correct grammar, and identify main sections with headings. Ensure a professional tone."
    model: str = "claude-3-5-haiku-20241022"
    max_tokens: int = 4000

# Pydantic model for Gemini User queries (now available for all paid users)
class UserAIGeminiRequest_Pydantic(BaseModel):
    transcript: str
    user_prompt: str
    model: str = "models/gemini-pro-latest"
    max_tokens: int = 1000

# Pydantic model for Gemini Admin formatting
class AdminAIFormatGeminiRequest_Pydantic(BaseModel):
    transcript: str
    formatting_instructions: str = "Correct all grammar, ensure a formal tone, break into paragraphs with subheadings for each major topic, and highlight action items in bold."
    model: str = "models/gemini-pro-latest"
    max_tokens: int = 4000

jobs = {}
active_background_tasks = {}
cancellation_flags = {}

logger.info("Enhanced job tracking initialized")

# Firebase Firestore interaction functions
async def update_user_plan_firestore(user_id: str, new_plan: str, reference_id: Optional[str] = None, payment_amount_usd: Optional[float] = None):
    """Updates a user's plan and related fields in Firestore using Firebase Admin SDK."""
    if not db:
        logger.error("Firestore client not initialized. Cannot update user plan.")
        return {'success': False, 'error': 'Firestore not initialized'}

    user_ref = db.collection('users').document(user_id)
    updates = {
        'plan': new_plan,
        'lastAccessed': firestore.SERVER_TIMESTAMP,
        'paystackReferenceId': reference_id,
        'hasReceivedInitialFreeMinutes': True,
        'totalMinutesUsed': 0
    }

    plan_duration_days = 0
    if new_plan == 'Three-Day Plan':
        plan_duration_days = 3
    elif new_plan == 'One-Week Plan':
        plan_duration_days = 7
    elif new_plan == 'Monthly Plan':
        plan_duration_days = 30
    elif new_plan == 'Yearly Plan':
        plan_duration_days = 365

    if plan_duration_days > 0:
        expires_at = datetime.now() + timedelta(days=plan_duration_days)
        updates['expiresAt'] = expires_at
        updates['subscriptionStartDate'] = firestore.SERVER_TIMESTAMP
        logger.info(f"User {user_id} {new_plan} plan will expire on: {expires_at}")
    else:
        updates['expiresAt'] = None
        updates['subscriptionStartDate'] = None

    try:
        await asyncio.to_thread(user_ref.update, updates) 
        logger.info(f"User {user_id} plan updated to {new_plan} in Firestore.")
        return {'success': True}
    except Exception as e:
        logger.error(f"Error updating user {user_id} plan in Firestore: {e}")
        return {'success': False, 'error': str(e)}

async def update_monthly_revenue_firestore(amount: float):
    """Updates the cumulative monthly revenue in Firestore."""
    if not db:
        logger.error("Firestore client not initialized. Cannot update monthly revenue.")
        return {'success': False, 'error': 'Firestore not initialized'}

    admin_stats_ref = db.collection('admin_stats').document('current')
    try:
        await asyncio.to_thread(admin_stats_ref.update, {
            'monthlyRevenue': firestore.Increment(amount)
        })
        logger.info(f"Monthly revenue updated by {amount} in Firestore.")
        return {'success': True}
    except Exception as e:
        logger.error(f"Error updating monthly revenue in Firestore: {e}")
        return {'success': False, 'error': str(e)}

async def get_user_profile_by_email_firestore(email: str):
    """Fetches user profile by email to get UID (for webhook processing)."""
    if not db:
        logger.error("Firestore client not initialized. Cannot fetch user by email.")
        return None
    try:
        users_ref = db.collection('users')
        query_ref = users_ref.where(filter=FieldFilter("email", "==", email)).limit(1)
        snapshot = await asyncio.to_thread(query_ref.get) 

        for doc in snapshot:
            return doc.id
        return None
    except Exception as e:
        logger.error(f"Error fetching user by email {email}: {e}")
        return None

async def analyze_audio_characteristics(audio_path: str) -> dict:
    try:
        audio = AudioSegment.from_file(audio_path)
        duration_seconds = len(audio) / 1000.0
        
        if audio.dBFS < -50:
            quality_score = 0.1
        elif audio.dBFS < -30:
            quality_score = 0.4
        else:
            quality_score = 0.8
            
        language = "unknown"
        return {
            "duration_seconds": duration_seconds,
            "quality_score": quality_score,
            "language": language,
            "channels": audio.channels,
            "sample_rate": audio.frame_rate,
            "size_mb": os.path.getsize(audio_path) / (1024 * 1024)
        }
    except Exception as e:
        logger.error(f"Error analyzing audio characteristics: {e}")
        return {
            "duration_seconds": 0,
            "quality_score": 0,
            "language": "unknown",
            "channels": 0,
            "sample_rate": 0,
            "size_mb": 0,
            "error": str(e)
        }

def compress_audio_for_transcription(input_path: str, output_path: str = None, job_id: str = None) -> tuple[str, dict]:
    """Compress audio file optimally for transcription with cancellation support"""
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}_compressed.mp3"
    
    try:
        if job_id and cancellation_flags.get(job_id, False):
            logger.info(f"Job {job_id} cancelled during compression setup")
            raise asyncio.CancelledError(f"Job {job_id} was cancelled")
            
        logger.info(f"Compressing {input_path} for transcription...")
        
        input_size = os.path.getsize(input_path) / (1024 * 1024)
        logger.info(f"Original file size: {input_size:.2f} MB")
        
        audio = AudioSegment.from_file(input_path)
        logger.info(f"Original audio: {audio.channels} channels, {audio.frame_rate}Hz, {len(audio)}ms")
        
        if job_id and cancellation_flags.get(job_id, False):
            logger.info(f"Job {job_id} cancelled during audio loading")
            raise asyncio.CancelledError(f"Job {job_id} was cancelled")
        
        if audio.channels > 1:
            audio = audio.set_channels(1)
            logger.info("Converted to mono audio")
        
        target_sample_rate = 16000
        audio = audio.set_frame_rate(target_sample_rate)
        logger.info(f"Reduced sample rate to {target_sample_rate} Hz")
        
        if job_id and cancellation_flags.get(job_id, False):
            logger.info(f"Job {job_id} cancelled during sample rate conversion")
            raise asyncio.CancelledError(f"Job {job_id} was cancelled")
        
        audio = audio - 3
        audio = audio.normalize()
        logger.info("Applied audio normalization")
        
        if job_id and cancellation_flags.get(job_id, False):
            logger.info(f"Job {job_id} cancelled before export")
            raise asyncio.CancelledError(f"Job {job_id} was cancelled")
        
        audio.export(
            output_path, 
            format="mp3",
            bitrate="64k",
            parameters=[
                "-q:a", "9",
                "-ac", "1",
                "-ar", str(target_sample_rate)
            ]
        )
        logger.info("Audio compression complete")
        
        if os.path.exists(output_path):
            output_size = os.path.getsize(output_path) / (1024 * 1024)
            
            size_difference = input_size - output_size
            if input_size > 0:
                compression_ratio = (size_difference / input_size) * 100
            else:
                compression_ratio = 0
            
            stats = {
                "original_size_mb": round(input_size, 2),
                "compressed_size_mb": round(output_size, 2),
                "compression_ratio_percent": round(compression_ratio, 1),
                "size_reduction_mb": round(size_difference, 2),
                "duration_seconds": len(audio) / 1000.0
            }
            
            logger.info(f"Compression result:")
            logger.info(f"  Original: {stats['original_size_mb']} MB")
            logger.info(f"  Processed: {stats['compressed_size_mb']} MB")
            if size_difference > 0:
                logger.info(f"  Size reduction: {stats['compression_ratio_percent']}% ({stats['size_reduction_mb']} MB saved)")
            else:
                logger.info(f"  Size increase: {abs(stats['compression_ratio_percent'])}% ({abs(stats['size_reduction_mb'])} MB added)")
        
        return output_path, stats
        
    except asyncio.CancelledError:
        logger.info(f"Compression cancelled for job {job_id}")
        if os.path.exists(output_path):
            os.unlink(output_path)
        raise
        
    except Exception as e:
        logger.error(f"Error compressing audio: {e}")
        try:
            if job_id and cancellation_flags.get(job_id, False):
                raise asyncio.CancelledError(f"Job {job_id} was cancelled")
                
            audio = AudioSegment.from_file(input_path)
            audio = audio.set_channels(1)
            audio = audio.set_frame_rate(16000)
            audio.export(output_path, format="mp3", bitrate="64k")
            
            output_size = os.path.getsize(output_path) / (1024 * 1024)
            size_difference = input_size - output_size
            compression_ratio = (size_difference / input_size) * 100 if input_size > 0 else 0
            
            stats = {
                "original_size_mb": round(input_size, 2),
                "compressed_size_mb": round(output_size, 2),
                "compression_ratio_percent": round(compression_ratio, 1),
                "size_reduction_mb": round(size_difference, 2),
                "duration_seconds": len(audio) / 1000.0
            }
            
            logger.info("Used fallback compression")
            return output_path, stats
            
        except asyncio.CancelledError:
            logger.info(f"Fallback compression cancelled for job {job_id}")
            if os.path.exists(output_path):
                os.unlink(output_path)
            raise
            
        except Exception as fallback_error:
            logger.error(f"Fallback compression also failed: {fallback_error}")
            raise
def compress_audio_for_download(input_path: str, output_path: str = None, quality: str = "high") -> str:
    """Compress audio file for download with different quality options"""
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}_download.mp3"
    
    try:
        logger.info(f"Compressing {input_path} for download (quality: {quality})...")
        
        audio = AudioSegment.from_file(input_path)
        
        if quality == "high":
            bitrate = "128k"
            sample_rate = 44100
            channels = 2 if audio.channels > 1 else 1
        elif quality == "medium":
            bitrate = "96k"
            sample_rate = 22050
            channels = 1
        else:
            bitrate = "64k"
            sample_rate = 16000
            channels = 1
        
        if audio.channels != channels:
            audio = audio.set_channels(channels)
        if audio.frame_rate != sample_rate:
            audio = audio.set_frame_rate(sample_rate)
        
        audio.export(
            output_path,
            format="mp3",
            bitrate=bitrate,
            parameters=[
                "-q:a", "2" if quality == "high" else "5",
                "-ac", str(channels),
                "-ar", str(sample_rate)
            ]
        )
        
        logger.info(f"Download compression complete: {quality} quality")
        return output_path
        
    except Exception as e:
        logger.error(f"Error compressing audio for download: {e}")
        raise
# Currency Conversion and Channel Mapping Logic
USD_TO_LOCAL_RATES = {
    'KE': 145.0,
    'NG': 1500.0,
    'GH': 15.0,
    'ZA': 19.0,
    'OTHER_AFRICA': 'USD', # 'USD' indicates no conversion, direct USD payment
}

COUNTRY_CURRENCY_MAP = {
    'KE': 'KES',
    'NG': 'NGN',
    'GH': 'GHS',
    'ZA': 'ZAR',
    'OTHER_AFRICA': 'USD',
}

COUNTRY_CHANNELS_MAP = {
    'KE': ['mobile_money', 'card'],
    'NG': ['bank', 'ussd', 'mobile_money', 'card'],
    'GH': ['mobile_money', 'card'],
    'ZA': ['eft', 'card'],
    'OTHER_AFRICA': ['card'],
}

def get_local_amount_and_currency(base_usd_amount: float, country_code: str, plan_name: str = None) -> tuple[float, str]:
    # REMOVED: 'Monthly Plan' from this condition. Now only Yearly Plan forces USD.
    if plan_name in ['Yearly Plan']:
        return base_usd_amount, 'USD'

    currency = COUNTRY_CURRENCY_MAP.get(country_code, 'USD')
    if currency == 'USD':
        return base_usd_amount, 'USD'
    
    rate = USD_TO_LOCAL_RATES.get(country_code, 1.0)
    local_amount = round(base_usd_amount * rate, 2)
    return local_amount, currency

def get_payment_channels(country_code: str, plan_name: str = None) -> list[str]:
    # REMOVED: 'Monthly Plan' from this condition. Now only Yearly Plan forces card payments.
    if plan_name in ['Yearly Plan']:
        return ['card']
    return COUNTRY_CHANNELS_MAP.get(country_code, ['card'])

async def health_monitor():
    logger.info("Starting health monitor background task")
    while True:
        try:
            import psutil
            memory_info = psutil.virtual_memory()
            cpu_percent = psutil.cpu_percent(interval=1)
            logger.info(f"Health Check - Memory: {memory_info.percent}% used, CPU: {cpu_percent}%, Available RAM: {memory_info.available / (1024**3):.2f} GB")
            logger.info(f"Active jobs: {len(jobs)}, Active background tasks: {len(active_background_tasks)}, Cancellation flags: {len(cancellation_flags)}")
            await asyncio.sleep(30)
        except Exception as e:
            logger.error(f"Health monitor error: {e}")
            await asyncio.sleep(30)

async def verify_paystack_payment(reference: str) -> dict:
    """Verify Paystack payment using reference"""
    if not PAYSTACK_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Paystack configuration missing")
    
    headers = {
        'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}',
        'Content-Type': 'application/json'
    }
    
    try:
        logger.info(f"Verifying Paystack payment with reference: {reference}")
        response = requests.get(
            f'https://api.paystack.co/transaction/verify/{reference}',
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            payment_data = response.json()
            
            if payment_data['status'] and payment_data['data']['status'] == 'success':
                amount_kobo = payment_data['data']['amount']
                amount = amount_kobo / 100
                customer_email = payment_data['data']['customer']['email']
                currency = payment_data['data']['data']['currency'] # Corrected: access currency from data.data
                plan_name = payment_data['data']['metadata'].get('plan', 'Unknown')
                
                logger.info(f"✅ Paystack payment verified: {customer_email} paid {amount} {currency} for {plan_name}")
                
                return {
                    'status': 'success',
                    'amount': amount,
                    'currency': currency,
                    'email': customer_email,
                    'plan': plan_name,
                    'reference': reference,
                    'raw_data': payment_data['data']
                }
            else:
                logger.warning(f"❌ Paystack payment verification failed: {payment_data}")
                return {
                    'status': 'failed',
                    'error': payment_data.get('message', 'Payment verification failed'),
                    'raw_data': payment_data
                }
        else:
            logger.error(f"❌ Paystack API error: {response.status_code} - {response.text}")
            return {
                'status': 'error',
                'error': f'Paystack API error: {response.status_code}',
                'details': response.text
            }
            
    except requests.exceptions.RequestException as e:
        logger.error(f"❌ Network error during Paystack verification: {str(e)}")
        return {
            'status': 'error',
            'error': 'Network error during payment verification',
            'details': str(e)
        }
    except Exception as e:
        logger.error(f"❌ Unexpected error during Paystack verification: {str(e)}")
        return {
            'status': 'error',
            "error": 'Payment verification failed',
            'details': str(e)
        }

async def update_user_credits_paystack(email: str, plan_name: str, amount: float, currency: str, update_admin_revenue: bool = False, country_code: Optional[str] = None):
    """
    Update user credits/plan in Firestore.
    The real-time revenue counter logic is now handled purely on the frontend.
    """
    if not db:
        logger.error(f"Firestore client not initialized. Cannot update credits for {email}.")
        return {'success': False, 'error': 'Firestore not initialized'}

    try:
        logger.info(f"📝 Updating credits for {email} - {plan_name} ({amount} {currency}) in Firestore.")
        
        # 1. Get user UID from email
        user_id = await get_user_profile_by_email_firestore(email)
        if not user_id:
            logger.error(f"User with email {email} not found in Firestore. Cannot update plan.")
            return {'success': False, 'error': f"User {email} not found in Firestore."}

        # 2. Update user's plan in Firestore
        await asyncio.to_thread(db.collection('users').document(user_id).update, {
            'plan': plan_name,
            'lastAccessed': firestore.SERVER_TIMESTAMP,
            'paystackReferenceId': None, 
            'hasReceivedInitialFreeMinutes': True,
            'totalMinutesUsed': 0
        })

        # 3. Update monthly revenue if flag is True
        if update_admin_revenue:
            revenue_update_result = await update_monthly_revenue_firestore(amount)
            if revenue_update_result['success']:
                logger.info(f"✅ Monthly revenue updated successfully for payment of {amount} {currency}.")
            else:
                logger.warning(f"⚠️ Failed to update monthly revenue: {revenue_update_result.get('error')}")
        
        logger.info(f"✅ Credits and plan updated successfully for {email} in Firestore.")
        return {'success': True, 'email': email, 'plan': plan_name, 'amount': amount, 'currency': currency}
        
    except Exception as e:
        logger.error(f"❌ Error updating user credits in Firestore: {str(e)}")
        return {'success': False, 'error': str(e)}

async def transcribe_with_openai_whisper(audio_path: str, language_code: str, job_id: str) -> dict:
    """Calls the dedicated OpenAI Whisper service deployed on Render."""
    if not OPENAI_WHISPER_SERVICE_RAILWAY_URL:
        logger.error(f"{TYPEMYWORDZ2_NAME} Service URL not configured, skipping {TYPEMYWORDZ2_NAME} for job {job_id}")
        return {
            "status": "failed",
            "error": "Transcription service unavailable. Please try again later."
        }

    try:
        logger.info(f"Calling {TYPEMYWORDZ2_NAME} service for job {job_id} at {OPENAI_WHISPER_SERVICE_RAILWAY_URL}/transcribe")
        
        # Read the audio file content
        with open(audio_path, "rb") as f:
            audio_content = f.read()

        # Prepare form data
        files = {'file': (os.path.basename(audio_path), audio_content, 'audio/mpeg')}
        data = {'language_code': language_code}

        # Make HTTP POST request to the dedicated Whisper service
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{OPENAI_WHISPER_SERVICE_RAILWAY_URL}/transcribe",
                files=files,
                data=data,
                timeout=300.0
            )
            response.raise_for_status()

        result = response.json()
        
        if result.get("status") == "completed" and result.get("transcription"):
            logger.info(f"{TYPEMYWORDZ2_NAME} service transcription completed for job {job_id}")
            return result
        else:
            raise Exception(f"{TYPEMYWORDZ2_NAME} service returned an incomplete or failed status: {result}")

    except asyncio.CancelledError:
        logger.info(f"{TYPEMYWORDZ2_NAME} service call cancelled for job {job_id}")
        raise
    except httpx.HTTPStatusError as e:
        logger.error(f"{TYPEMYWORDZ2_NAME} service HTTP error for job {job_id}: {e.response.status_code} - {e.response.text}")
        return {
            "status": "failed",
            "error": "Transcription failed due to a service error. Please try again."
        }
    except httpx.RequestError as e:
        logger.error(f"{TYPEMYWORDZ2_NAME} service network error for job {job_id}: {e}")
        return {
            "status": "failed",
            "error": "Transcription failed due to a network issue. Please check your connection and try again."
        }
    except Exception as e:
        logger.error(f"{TYPEMYWORDZ2_NAME} transcription failed for job {job_id}: {str(e)}")
        return {
            "status": "failed",
            "error": "Transcription failed. Please try again later."
        }
    finally:
        pass

async def transcribe_with_assemblyai(audio_path: str, language_code: str, speaker_labels_enabled: bool, model: str, job_id: str) -> dict:
    """Transcribe audio using AssemblyAI API"""
    if not ASSEMBLYAI_API_KEY:
        logger.error(f"{TYPEMYWORDZ1_NAME} API Key not configured, skipping {TYPEMYWORDZ1_NAME} for job {job_id}")
        return {
            "status": "failed",
            "error": "Transcription service unavailable. Please try again later."
        }

    try:
        logger.info(f"Starting {TYPEMYWORDZ1_NAME} transcription with {model} model for job {job_id}")
        
        def check_cancellation():
            if job_id and cancellation_flags.get(job_id, False):
                logger.info(f"Job {job_id} was cancelled during {TYPEMYWORDZ1_NAME} processing")
                raise asyncio.CancelledError(f"Job {job_id} was cancelled")
        
        check_cancellation()
        
        compressed_path, compression_stats = compress_audio_for_transcription(audio_path, job_id=job_id)
        logger.info(f"Audio compressed for {TYPEMYWORDZ1_NAME}: {compression_stats}")

        check_cancellation()
        
        logger.info(f"Uploading audio to {TYPEMYWORDZ1_NAME}...")
        headers = {"authorization": ASSEMBLYAI_API_KEY}
        upload_endpoint = "https://api.assemblyai.com/v2/upload"
        
        with open(compressed_path, "rb") as f:
            upload_response = requests.post(upload_endpoint, headers=headers, data=f)
        
        if upload_response.status_code != 200:
            raise Exception(f"Failed to upload audio to {TYPEMYWORDZ1_NAME}: {upload_response.text}")
        
        upload_result = upload_response.json()
        audio_url = upload_result["upload_url"]
        logger.info(f"Audio uploaded to {TYPEMYWORDZ1_NAME}: {audio_url}")

        check_cancellation()

        headers = {"authorization": ASSEMBLYAI_API_KEY, "content-type": "application/json"}
        transcript_endpoint = "https://api.assemblyai.com/v2/transcript"
        json_data = {
            "audio_url": audio_url,
            "language_code": language_code,
            "punctuate": True,
            "format_text": True,
            "speaker_labels": speaker_labels_enabled,
            "speech_model": model,
            "word_boost": []
        }
        
        transcript_response = requests.post(transcript_endpoint, headers=headers, json=json_data)
        
        if transcript_response.status_code != 200:
            raise Exception(f"Failed to start transcription on {TYPEMYWORDZ1_NAME}: {transcript_response.text}")
        
        transcript_result = transcript_response.json()
        transcript_id = transcript_result["id"]
        logger.info(f"{TYPEMYWORDZ1_NAME} transcription started with ID: {transcript_id}")
        
        if os.path.exists(compressed_path):
            os.unlink(compressed_path)
            logger.info(f"Cleaned up compressed file: {compressed_path}")

        while True:
            check_cancellation()
            
            await asyncio.sleep(5)
            
            status_response = requests.get(f"https://api.assemblyai.com/v2/transcript/{transcript_id}", headers={"authorization": ASSEMBLYAI_API_KEY})
            
            if status_response.status_code != 200:
                raise Exception(f"Failed to get status from {TYPEMYWORDZ1_NAME}: {status_response.text}")
            
            status_result = status_response.json()
            
            if status_result["status"] == "completed":
                transcription_text = status_result["text"]
                
                if speaker_labels_enabled and status_result.get("utterances"):
                    formatted_transcript = ""
                    for utterance in status_result.get("utterances"):
                        speaker_letter = utterance['speaker']
                        if speaker_letter == 'A':
                            speaker_num = '1'
                        elif speaker_letter == 'B':
                            speaker_num = '2'
                        elif speaker_letter == 'C':
                            speaker_num = '3'
                        elif speaker_letter == 'D':
                            speaker_num = '4'
                        elif speaker_letter == 'E':
                            speaker_num = '5'
                        else:
                            speaker_num = str(ord(speaker_letter.upper()) - ord('A') + 1)
                        
                        formatted_transcript += f"<strong>Speaker {speaker_num}:</strong> {utterance['text']}\n"
                    transcription_text = formatted_transcript.strip()

                return {
                    "status": "completed",
                    "transcription": transcription_text,
                    "language": status_result["language_code"],
                    "duration": status_result.get("audio_duration", 0),
                    "word_count": len(transcription_text.split()) if transcription_text else 0,
                    "has_speaker_labels": speaker_labels_enabled and bool(status_result.get("utterances"))
                }
            elif status_result["status"] == "error":
                raise Exception(status_result.get("error", f"Transcription failed on {TYPEMYWORDZ1_NAME}"))
            else:
                logger.info(f"{TYPEMYWORDZ1_NAME} status: {status_result['status']}")
                continue
        
    except asyncio.CancelledError:
        logger.info(f"{TYPEMYWORDZ1_NAME} transcription cancelled for job {job_id}")
        raise
    except Exception as e:
        logger.error(f"{TYPEMYWORDZ1_NAME} transcription failed for job {job_id}: {str(e)}")
        return {
            "status": "failed",
            "error": "Transcription failed. Please try again later."
        }

async def transcribe_with_deepgram(audio_path: str, language_code: str, speaker_labels_enabled: bool, job_id: str) -> dict:
    """Calls the dedicated Deepgram service deployed on Render."""
    if not DEEPGRAM_SERVICE_RAILWAY_URL:
        logger.error(f"{DEEPGRAM_NAME} Service URL not configured, skipping {DEEPGRAM_NAME} for job {job_id}")
        return {
            "status": "failed",
            "error": "Transcription service unavailable. Please try again later."
        }

    try:
        logger.info(f"Calling {DEEPGRAM_NAME} service for job {job_id} at {DEEPGRAM_SERVICE_RAILWAY_URL}/transcribe")
        
        # Read the audio file content
        with open(audio_path, "rb") as f:
            audio_content = f.read()

        # Prepare form data
        files = {'file': (os.path.basename(audio_path), audio_content, 'audio/mpeg')}
        data = {
            'language_code': language_code,
            'speaker_labels_enabled': str(speaker_labels_enabled).lower()
        }

        # Make HTTP POST request to the dedicated Deepgram service
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{DEEPGRAM_SERVICE_RAILWAY_URL}/transcribe",
                files=files,
                data=data,
                timeout=300.0
            )
            response.raise_for_status()

        result = response.json()
        
        if result.get("status") == "completed" and result.get("transcription"):
            logger.info(f"{DEEPGRAM_NAME} service transcription completed for job {job_id}")
            return result
        else:
            raise Exception(f"{DEEPGRAM_NAME} service returned an incomplete or failed status: {result}")

    except asyncio.CancelledError:
        logger.info(f"{DEEPGRAM_NAME} service call cancelled for job {job_id}")
        raise
    except httpx.HTTPStatusError as e:
        logger.error(f"{DEEPGRAM_NAME} service HTTP error for job {job_id}: {e.response.status_code} - {e.response.text}")
        return {
            "status": "failed",
            "error": "Transcription failed due to a service error. Please try again."
        }
    except httpx.RequestError as e:
        logger.error(f"{DEEPGRAM_NAME} service network error for job {job_id}: {e}")
        return {
            "status": "failed",
            "error": "Transcription failed due to a network issue. Please check your connection and try again."
        }
    except Exception as e:
        logger.error(f"{DEEPGRAM_NAME} transcription failed for job {job_id}: {str(e)}")
        return {
            "status": "failed",
            "error": "Transcription failed. Please try again later."
        }
    finally:
        pass

async def process_transcription_job(job_id: str, tmp_path: str, filename: str, language_code: Optional[str], speaker_labels_enabled: bool, user_plan: str, duration_minutes: float, user_email: str = ""):
    """Updated transcription processing with new three-tier service logic and admin/tester email checking."""
    logger.info(f"Starting transcription job {job_id}: {filename}, duration: {duration_minutes:.1f}min, plan: {user_plan}, email: {user_email}, speaker_labels: {speaker_labels_enabled}")
    job_data = jobs[job_id]
    
    active_background_tasks[job_id] = asyncio.current_task()
    cancellation_flags[job_id] = False

    compressed_path = None

    try:
        def check_cancellation():
            if cancellation_flags.get(job_id, False) or job_data.get("status") == "cancelled":
                logger.info(f"Job {job_id} was cancelled - stopping processing")
                raise asyncio.CancelledError(f"Job {job_id} was cancelled")
            return True

        check_cancellation()

        # Get service configuration based on new logic
        service_config = get_transcription_services(user_plan, speaker_labels_enabled, user_email)
        tier_1_service = service_config["tier_1"]
        tier_2_service = service_config["tier_2"]
        tier_3_service = service_config["tier_3"]

        logger.info(f"🎯 Job {job_id} service selection: Tier1={tier_1_service}, Tier2={tier_2_service}, Tier3={tier_3_service} ({service_config['reason']})")

        def get_assemblyai_model(plan: str) -> str:
            if plan == 'free' or plan == 'Monthly Plan':
                return "nano"  
            else:
                return "best"  

        assemblyai_model = get_assemblyai_model(user_plan)

        job_data.update({
            "tier_1_service": tier_1_service,
            "tier_2_service": tier_2_service,
            "tier_3_service": tier_3_service,
            "assemblyai_model": assemblyai_model,
            "duration_minutes": duration_minutes,
            "selection_reason": service_config["reason"],
            "user_email": user_email,
            "is_admin": is_admin_user(user_email)
        })

        transcription_result = None
        services_attempted = []

        # --- ATTEMPT TIER 1 SERVICE ---
        if tier_1_service == "assemblyai":
            if not ASSEMBLYAI_API_KEY:
                logger.error(f"{TYPEMYWORDZ1_NAME} API Key not configured, skipping Tier 1 for job {job_id}")
                job_data["tier_1_error"] = f"{TYPEMYWORDZ1_NAME} API Key not configured"
            else:
                try:
                    logger.info(f"🚀 Attempting {TYPEMYWORDZ1_NAME} (Tier 1 Primary) for job {job_id}")
                    compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                    logger.info(f"Audio compressed for {TYPEMYWORDZ1_NAME}: {compression_stats}")
                    transcription_result = await transcribe_with_assemblyai(compressed_path, language_code, speaker_labels_enabled, assemblyai_model, job_id)
                    job_data["tier_1_used"] = "assemblyai"
                    job_data["tier_1_success"] = True
                except Exception as error:
                    logger.error(f"❌ {TYPEMYWORDZ1_NAME} (Tier 1 Primary) failed: {error}")
                    job_data["tier_1_error"] = str(error)
                    job_data["tier_1_success"] = False
            services_attempted.append(f"{TYPEMYWORDZ1_NAME}_tier1")
            
        elif tier_1_service == "openai_whisper":
            if not OPENAI_WHISPER_SERVICE_RAILWAY_URL:
                logger.error(f"{TYPEMYWORDZ2_NAME} Service URL not configured, skipping Tier 1 for job {job_id}")
                job_data["tier_1_error"] = f"{TYPEMYWORDZ2_NAME} Service URL not configured"
            else:
                try:
                    logger.info(f"🚀 Attempting {TYPEMYWORDZ2_NAME} (Tier 1 Primary) for job {job_id}")
                    compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                    logger.info(f"Audio compressed for {TYPEMYWORDZ2_NAME}: {compression_stats}")
                    transcription_result = await transcribe_with_openai_whisper(compressed_path, language_code, job_id)
                    job_data["tier_1_used"] = "openai_whisper"
                    job_data["tier_1_success"] = True
                except Exception as error:
                    logger.error(f"❌ {TYPEMYWORDZ2_NAME} (Tier 1 Primary) failed: {error}")
                    job_data["tier_1_error"] = str(error)
                    job_data["tier_1_success"] = False
            services_attempted.append(f"{TYPEMYWORDZ2_NAME}_tier1")
        
        elif tier_1_service == "deepgram":
            if not DEEPGRAM_SERVICE_RAILWAY_URL:
                logger.error(f"{DEEPGRAM_NAME} Service URL not configured, skipping Tier 1 for job {job_id}")
                job_data["tier_1_error"] = f"{DEEPGRAM_NAME} Service URL not configured"
            else:
                try:
                    logger.info(f"🚀 Attempting {DEEPGRAM_NAME} (Tier 1 Primary) for job {job_id}")
                    compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                    logger.info(f"Audio compressed for {DEEPGRAM_NAME}: {compression_stats}")
                    transcription_result = await transcribe_with_deepgram(compressed_path, language_code, speaker_labels_enabled, job_id)
                    job_data["tier_1_used"] = "deepgram"
                    job_data["tier_1_success"] = True
                except Exception as error:
                    logger.error(f"❌ {DEEPGRAM_NAME} (Tier 1 Primary) failed: {error}")
                    job_data["tier_1_error"] = str(error)
                    job_data["tier_1_success"] = False
            services_attempted.append(f"{DEEPGRAM_NAME}_tier1")
        
        check_cancellation()

        # --- ATTEMPT TIER 2 SERVICE (FALLBACK 1) if Tier 1 failed AND tier_2_service is defined ---
        if (not transcription_result or transcription_result.get("status") == "failed") and tier_2_service:
            logger.warning(f"⚠️ Tier 1 service failed, trying Tier 2 fallback ({tier_2_service}) for job {job_id}")
            
            if tier_2_service == "assemblyai":
                if not ASSEMBLYAI_API_KEY:
                    logger.error(f"{TYPEMYWORDZ1_NAME} API Key not configured, skipping Tier 2 for job {job_id}")
                    job_data["tier_2_error"] = f"{TYPEMYWORDZ1_NAME} API Key not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {TYPEMYWORDZ1_NAME} (Tier 2 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {TYPEMYWORDZ1_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_assemblyai(compressed_path, language_code, speaker_labels_enabled, assemblyai_model, job_id)
                        job_data["tier_2_used"] = "assemblyai"
                        job_data["tier_2_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {TYPEMYWORDZ1_NAME} (Tier 2 Fallback) failed: {error}")
                        job_data["tier_2_error"] = str(error)
                        job_data["tier_2_success"] = False
                services_attempted.append(f"{TYPEMYWORDZ1_NAME}_tier2")
                
            elif tier_2_service == "openai_whisper":
                if not OPENAI_WHISPER_SERVICE_RAILWAY_URL:
                    logger.error(f"{TYPEMYWORDZ2_NAME} Service URL not configured, skipping Tier 2 for job {job_id}")
                    job_data["tier_2_error"] = f"{TYPEMYWORDZ2_NAME} Service URL not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {TYPEMYWORDZ2_NAME} (Tier 2 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {TYPEMYWORDZ2_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_openai_whisper(compressed_path, language_code, job_id)
                        job_data["tier_2_used"] = "openai_whisper"
                        job_data["tier_2_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {TYPEMYWORDZ2_NAME} (Tier 2 Fallback) failed: {error}")
                        job_data["tier_2_error"] = str(error)
                        job_data["tier_2_success"] = False
                services_attempted.append(f"{TYPEMYWORDZ2_NAME}_tier2")
            
            elif tier_2_service == "deepgram":
                if not DEEPGRAM_SERVICE_RAILWAY_URL:
                    logger.error(f"{DEEPGRAM_NAME} Service URL not configured, skipping Tier 2 for job {job_id}")
                    job_data["tier_2_error"] = f"{DEEPGRAM_NAME} Service URL not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {DEEPGRAM_NAME} (Tier 2 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {DEEPGRAM_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_deepgram(compressed_path, language_code, speaker_labels_enabled, job_id)
                        job_data["tier_2_used"] = "deepgram"
                        job_data["tier_2_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {DEEPGRAM_NAME} (Tier 2 Fallback) failed: {error}")
                        job_data["tier_2_error"] = str(error)
                        job_data["tier_2_success"] = False
                services_attempted.append(f"{DEEPGRAM_NAME}_tier2")
            
        check_cancellation()

        # --- ATTEMPT TIER 3 SERVICE (FALLBACK 2) if Tier 2 failed AND tier_3_service is defined ---
        if (not transcription_result or transcription_result.get("status") == "failed") and tier_3_service:
            logger.warning(f"⚠️ Tier 2 service failed, trying Tier 3 fallback ({tier_3_service}) for job {job_id}")
            
            if tier_3_service == "assemblyai":
                if not ASSEMBLYAI_API_KEY:
                    logger.error(f"{TYPEMYWORDZ1_NAME} API Key not configured, skipping Tier 3 for job {job_id}")
                    job_data["tier_3_error"] = f"{TYPEMYWORDZ1_NAME} API Key not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {TYPEMYWORDZ1_NAME} (Tier 3 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {TYPEMYWORDZ1_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_assemblyai(compressed_path, language_code, speaker_labels_enabled, assemblyai_model, job_id)
                        job_data["tier_3_used"] = "assemblyai"
                        job_data["tier_3_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {TYPEMYWORDZ1_NAME} (Tier 3 Fallback) failed: {error}")
                        job_data["tier_3_error"] = str(error)
                        job_data["tier_3_success"] = False
                services_attempted.append(f"{TYPEMYWORDZ1_NAME}_tier3")
                
            elif tier_3_service == "openai_whisper":
                if not OPENAI_WHISPER_SERVICE_RAILWAY_URL:
                    logger.error(f"{TYPEMYWORDZ2_NAME} Service URL not configured, skipping Tier 3 for job {job_id}")
                    job_data["tier_3_error"] = f"{TYPEMYWORDZ2_NAME} Service URL not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {TYPEMYWORDZ2_NAME} (Tier 3 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {TYPEMYWORDZ2_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_openai_whisper(compressed_path, language_code, job_id)
                        job_data["tier_3_used"] = "openai_whisper"
                        job_data["tier_3_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {TYPEMYWORDZ2_NAME} (Tier 3 Fallback) failed: {error}")
                        job_data["tier_3_error"] = str(error)
                        job_data["tier_3_success"] = False
                services_attempted.append(f"{TYPEMYWORDZ2_NAME}_tier3")
            
            elif tier_3_service == "deepgram":
                if not DEEPGRAM_SERVICE_RAILWAY_URL:
                    logger.error(f"{DEEPGRAM_NAME} Service URL not configured, skipping Tier 3 for job {job_id}")
                    job_data["tier_3_error"] = f"{DEEPGRAM_NAME} Service URL not configured"
                else:
                    try:
                        logger.info(f"🔄 Attempting {DEEPGRAM_NAME} (Tier 3 Fallback) for job {job_id}")
                        if compressed_path is None:
                            compressed_path, compression_stats = compress_audio_for_transcription(tmp_path, job_id=job_id)
                            logger.info(f"Audio compressed for {DEEPGRAM_NAME}: {compression_stats}")
                        transcription_result = await transcribe_with_deepgram(compressed_path, language_code, speaker_labels_enabled, job_id)
                        job_data["tier_3_used"] = "deepgram"
                        job_data["tier_3_success"] = True
                    except Exception as error:
                        logger.error(f"❌ {DEEPGRAM_NAME} (Tier 3 Fallback) failed: {error}")
                        job_data["tier_3_error"] = str(error)
                        job_data["tier_3_success"] = False
                services_attempted.append(f"{DEEPGRAM_NAME}_tier3")
            
        check_cancellation()

        # Final result processing
        if not transcription_result or transcription_result.get("status") == "failed":
            logger.error(f"❌ All transcription services failed for job {job_id}. Services attempted: {services_attempted}")
            job_data.update({
                "status": "failed",
                "error": "Transcription failed after multiple attempts. Please try again later.",
                "completed_at": datetime.now().isoformat(),
                "services_attempted": services_attempted
            })
        else:
            logger.info(f"✅ Transcription completed successfully for job {job_id}")
            service_used_name = (job_data.get("tier_1_used") or job_data.get("tier_2_used") or job_data.get("tier_3_used"))
            model_used = "N/A"

            if service_used_name == "assemblyai":
                model_used = job_data.get("assemblyai_model", "unknown")
            elif service_used_name == "openai_whisper":
                model_used = "whisper-1"
            elif service_used_name == "deepgram":
                model_used = "deepgram-nova"

            logger.info(f"📊 Job {job_id} for user {user_email} completed. Service: {service_used_name}, Model: {model_used}")

            job_data.update({
                "status": "completed",
                "transcription": transcription_result["transcription"],
                "language": transcription_result.get("language", language_code),
                "completed_at": datetime.now().isoformat(),
                "word_count": transcription_result.get("word_count", 0),
                "duration_seconds": transcription_result.get("duration", 0),
                "speaker_labels": speaker_labels_enabled,
                "service_used": service_used_name,
                "model_used": model_used,
                "services_attempted": services_attempted
            })

    except asyncio.CancelledError:
        logger.info(f"Transcription job {job_id} was cancelled")
        if job_data.get("status") != "cancelled":
            job_data.update({
                "status": "cancelled",
                "cancelled_at": datetime.now().isoformat(),
                "error": "Job was cancelled by user"
            })
        raise
        
    except Exception as e:
        logger.error(f"Transcription job: ERROR during processing for job {job_id}: {str(e)}")
        import traceback
        logger.error(f"Transcription job: Full traceback: {traceback.format_exc()}")
        job_data.update({
            "status": "failed",
            "error": "An unexpected error occurred during transcription. Please try again later.",
            "completed_at": datetime.now().isoformat()
        })
    finally:
        # Clean up original temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
            logger.info(f"Cleaned up original temp file: {tmp_path}")
        
        if job_id in active_background_tasks:
            del active_background_tasks[job_id]
        if job_id in cancellation_flags:
            del cancellation_flags[job_id]
            
        logger.info(f"Transcription job completed for job ID: {job_id}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Application lifespan startup")
    health_task = asyncio.create_task(health_monitor())
    logger.info("Health monitor task created")
    yield
    logger.info("Application lifespan shutdown")
    health_task.cancel()
    for job_id, task in active_background_tasks.items():
        if not task.done():
            logger.info(f"Cancelling background task for job {job_id}")
            cancellation_flags[job_id] = True
            task.cancel()
    jobs.clear()
    active_background_tasks.clear()
    cancellation_flags.clear()
    logger.info("All background tasks cancelled and cleanup complete")

logger.info("Creating FastAPI app...")
app = FastAPI(title=f"Enhanced Transcription Service with {TYPEMYWORDZ1_NAME}, {TYPEMYWORDZ2_NAME}, {DEEPGRAM_NAME}, {TYPEMYWORDZ_AI_NAME} & Google Gemini", lifespan=lifespan)
logger.info("FastAPI app created successfully")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
logger.info("CORS middleware configured successfully")
@app.get("/")
async def root():
    logger.info("Root endpoint called")
    return {
        "message": f"Enhanced Transcription Service with {TYPEMYWORDZ1_NAME}, {TYPEMYWORDZ2_NAME}, {DEEPGRAM_NAME}, {TYPEMYWORDZ_AI_NAME} & Google Gemini is running!",
        "features": [
            f"AssemblyAI integration with smart model selection",
            f"OpenAI Whisper integration for transcription",
            f"Deepgram integration for transcription",
            "Three-tier automatic fallback between services",
            "Paystack payment integration",
            f"Speaker diarization for AssemblyAI and Deepgram",
            "Language selection for transcription",
            f"User-driven AI features (summarization, Q&A, and bullet points) via TypeMyworDz AI (Anthropic)",
            f"Admin-driven AI formatting via TypeMyworDz AI (Anthropic) and Google Gemini",
            "Google Gemini integration for AI queries - NOW AVAILABLE FOR ALL PAID USERS"
        ],
        "logic": {
            "free_user_transcription": f"Primary={TYPEMYWORDZ1_NAME} → Fallback1={DEEPGRAM_NAME} → Fallback2=None",
            "three_day_plan_transcription": f"Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}",
            "one_week_plan_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
            "monthly_plan_transcription": f"Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}",
            "yearly_plan_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
            "admin_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
            "speaker_labels_transcription": f"Always use {TYPEMYWORDZ1_NAME} first → Fallback1={DEEPGRAM_NAME} → Fallback2=None",
            "assemblyai_tester_transcription": f"Always use {TYPEMYWORDZ1_NAME} (no fallback for {ASSEMBLYAI_TESTER_EMAIL})",
            "free_users_assemblyai_model": f"{TYPEMYWORDZ1_NAME} nano model",
            "paid_users_assemblyai_model": f"{TYPEMYWORDZ1_NAME} best model",
            "ai_features_access": "Only for Three-Day, One-Week, Monthly Plan, and Yearly Plan plans",
            "gemini_access": "NOW AVAILABLE FOR ALL PAID AI USERS (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)",
            "assemblyai": f"TypeMyworDz1 (AssemblyAI)",
            "openai_whisper": f"TypeMyworDz2 (OpenAI Whisper-1)",
            "deepgram": f"Deepgram",
            "anthropic_ai": f"TypeMyworDz AI (Anthropic Claude)",
            "google_gemini_ai": "Google Gemini - Available for ALL paid AI users",
            "admin_emails": ADMIN_EMAILS,
            "assemblyai_tester_email": ASSEMBLYAI_TESTER_EMAIL
        },
        "stats": {
            "active_jobs": len(jobs),
            "background_tasks": len(active_background_tasks),
            "cancellation_flags": len(cancellation_flags)
        }
    }

@app.post("/api/initialize-paystack-payment")
async def initialize_paystack_payment(request: PaystackInitializationRequest):
    logger.info(f"Initializing Paystack payment for {request.email} in {request.country_code}: Base USD {request.amount}")
    
    if not PAYSTACK_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Paystack configuration missing")
    
    try:
        # Determine local amount and currency, now Monthly Plan does NOT force USD
        local_amount, local_currency = get_local_amount_and_currency(request.amount, request.country_code, request.plan_name)
        payment_channels = get_payment_channels(request.country_code, request.plan_name)

        amount_kobo = int(local_amount * 100)
        
        headers = {
            'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}',
            'Content-Type': 'application/json'
        }
        
        payload = {
            'email': request.email,
            'amount': amount_kobo,
            'currency': local_currency,
            'callback_url': request.callback_url,
            'channels': payment_channels,
            'metadata': {
                'plan': request.plan_name,
                'base_usd_amount': request.amount,
                'country_code': request.country_code,
                'custom_fields': [
                    {
                        'display_name': "Plan Type",
                        'variable_name': "plan_type",
                        'value': request.plan_name
                    },
                    {
                        'display_name': "Country",
                        'variable_name': "country",
                        'value': request.country_code
                    },
                    {
                        'display_name': "Update Admin Revenue",
                        'variable_name': "update_admin_revenue",
                        'value': str(request.update_admin_revenue)
                    }
                ]
            }
        }
        
        logger.info(f"DEBUG: Paystack payload for {request.country_code}: Amount={local_amount} {local_currency}, Channels={payment_channels}")

        response = requests.post(
            'https://api.paystack.co/transaction/initialize',
            headers=headers,
            json=payload,
            timeout=10
        )
        
        if response.status_code == 200:
            result = response.json()
            logger.info(f"✅ Paystack payment initialized: {result['data']['reference']}")
            return {
                'status': True,
                'authorization_url': result['data']['authorization_url'],
                'reference': result['data']['reference'],
                'local_amount': local_amount,
                'local_currency': local_currency
            }
        else:
            logger.error(f"❌ Paystack API error: {response.status_code} - {response.text}")
            raise HTTPException(status_code=response.status_code, detail=f"Paystack API error: {response.text}")
            
    except Exception as e:
        import traceback
        logger.error(f"❌ Error initializing Paystack payment: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Payment initialization failed: {str(e)}")

@app.post("/api/verify-payment")
async def verify_payment(request: PaystackVerificationRequest):
    logger.info(f"Payment verification request for reference: {request.reference}")
    
    try:
        verification_result = await verify_paystack_payment(request.reference)
        
        if verification_result['status'] == 'success':
            email = verification_result['email']
            plan_name = verification_result['plan']
            amount = verification_result['amount']
            currency = verification_result['currency']
            reference = verification_result['reference']
            
            # Extract base_usd_amount and country_code from raw_data metadata
            base_usd_amount = verification_result['raw_data'].get('metadata', {}).get('base_usd_amount')
            country_code = verification_result['raw_data'].get('metadata', {}).get('country_code')
            update_admin_revenue_flag = verification_result['raw_data'].get('metadata', {}).get('update_admin_revenue', 'False').lower() == 'true'

            # Pass base_usd_amount, country_code, and update_admin_revenue_flag
            credit_result = await update_user_credits_paystack(email, plan_name, base_usd_amount or amount, currency, update_admin_revenue_flag, country_code) 
            
            if credit_result['success']:
                logger.info(f"✅ Payment verified and credits updated for {email}")
                return {
                    "status": "success",
                    "message": "Payment verified successfully",
                    "data": {
                        "amount": amount,
                        "currency": currency,
                        "email": email,
                        "plan": plan_name,
                        "reference": request.reference,
                        "credits_updated": True
                    }
                }
            else:
                logger.warning(f"⚠️ Payment verified but credit update failed for {email}: {credit_result.get('error')}")
                return {
                    "status": "partial_success",
                    "message": "Payment verified but credit update failed",
                    "data": {
                        "amount": amount,
                        "currency": currency,
                        "email": email,
                        "plan": plan_name,
                        "reference": request.reference,
                        "credits_updated": False,
                        "credit_error": credit_result.get('error')
                    }
                }
        else:
            logger.warning(f"❌ Payment verification failed for reference: {request.reference}")
            raise HTTPException(
                status_code=400, 
                detail=verification_result.get('error', 'Payment verification failed'),
                headers={"X-Error-Details": verification_result.get('details', '')}
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Unexpected error during payment verification: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail=f"Payment verification failed: {str(e)}"
        )

@app.post("/api/paystack-webhook")
async def paystack_webhook(request: Request):
    try:
        body = await request.body()
        signature = request.headers.get('x-paystack-signature')
        
        logger.info(f"Received Paystack webhook with signature: {bool(signature)}")
        
        if not body:
            logger.warning("Empty webhook body received")
            raise HTTPException(status_code=400, detail="Empty webhook body")
        
        try:
            webhook_data = json.loads(body.decode('utf-8'))
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in webhook: {e}")
            raise HTTPException(status_code=400, detail="Invalid JSON payload")
        
        # Optional: Verify webhook signature for production.
        # if PAYSTACK_WEBHOOK_SECRET:
        #     import hmac
        #     import hashlib
        #     expected_signature = hmac.new(PAYSTACK_WEBHOOK_SECRET.encode('utf-8'), body, hashlib.sha512).hexdigest()
        #     if not hmac.compare_digest(expected_signature, signature):
        #         logger.warning("❌ Webhook signature mismatch!")
        #         raise HTTPException(status_code=400, detail="Invalid webhook signature")
        
        event_type = webhook_data.get('event')
        logger.info(f"Processing Paystack webhook event: {event_type}")
        
        if event_type == 'charge.success':
            data = webhook_data.get('data', {})
            customer_email = data.get('customer', {}).get('email')
            amount = data.get('amount', 0) / 100
            currency = data.get('currency')
            reference = data.get('reference')
            plan_name = data.get('metadata', {}).get('plan', 'Unknown')
            base_usd_amount = data.get('metadata', {}).get('base_usd_amount')
            country_code = data.get('metadata', {}).get('country_code')
            update_admin_revenue_flag = data.get('metadata', {}).get('update_admin_revenue', 'False').lower() == 'true'

            logger.info(f"🔔 Webhook: Payment successful - {customer_email} paid {amount} {currency} for {plan_name}. Base USD: {base_usd_amount}, Country: {country_code}, Update Revenue: {update_admin_revenue_flag}")
            
            if customer_email:
                credit_result = await update_user_credits_paystack(customer_email, plan_name, base_usd_amount or amount, currency, update_admin_revenue_flag, country_code)
                if credit_result['success']:
                    logger.info(f"✅ Webhook: Credits updated automatically for {customer_email} in Firestore.")
                else:
                    logger.warning(f"⚠️ Webhook: Failed to update credits for {customer_email} in Firestore: {credit_result.get('error')}")
            
        elif event_type == 'charge.failed':
            data = webhook_data.get('data', {})
            customer_email = data.get('customer', {}).get('email')
            reference = data.get('reference')
            
            logger.warning(f"🔔 Webhook: Payment failed for {customer_email}, reference: {reference}")
            
        else:
            logger.info(f"🔔 Webhook: Unhandled event type: {event_type}")
        
        return {"status": "received", "event": event_type}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error processing Paystack webhook: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Webhook processing failed: {str(e)}")
@app.get("/api/paystack-status")
async def paystack_status():
    return {
        "paystack_configured": bool(PAYSTACK_SECRET_KEY),
        "public_key_configured": bool(PAYSTACK_PUBLIC_KEY),
        "webhook_secret_configured": bool(PAYSTACK_WEBHOOK_SECRET),
        "assemblyai_configured": bool(ASSEMBLYAI_API_KEY),
        "anthropic_configured": bool(ANTHROPIC_API_KEY),
        "openai_configured": bool(OPENAI_API_KEY),
        "openai_whisper_service_configured": bool(OPENAI_WHISPER_SERVICE_RAILWAY_URL),
        "google_gemini_configured": bool(GEMINI_API_KEY),
        "deepgram_service_configured": bool(DEEPGRAM_SERVICE_RAILWAY_URL),
        "admin_emails": ADMIN_EMAILS,
        "assemblyai_tester_email": ASSEMBLYAI_TESTER_EMAIL,
        "gemini_access": "NOW AVAILABLE FOR ALL PAID AI USERS (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)",
        "endpoints": {
            "initialize_payment": "/api/initialize-paystack-payment",
            "verify_payment": "/api/verify-payment",
            "webhook": "/api/paystack-webhook",
            "status": "/api/paystack-status",
            "transcribe": "/transcribe",
            "ai_user_query": "/ai/user-query",
            "ai_user_query_gemini": "/ai/user-query-gemini",
            "ai_admin_format": "/ai/admin-format",
            "ai_admin_format_gemini": "/ai/admin-format-gemini",
        },
        "supported_currencies": ["NGN", "USD", "GHS", "ZAR", "KES"],
        "supported_plans": [
            "Three-Day Plan",
            "One-Week Plan",
            "Monthly Plan",
            "Yearly Plan"
        ],
        "conversion_rates_usd_to_local": USD_TO_LOCAL_RATES
    }

@app.get("/api/list-gemini-models")
async def list_gemini_models():
    logger.info("Listing available Gemini models...")
    if not gemini_client:
        raise HTTPException(status_code=503, detail="Google Gemini service is not initialized (API key missing or invalid).")
    
    try:
        models = genai.list_models()
        gemini_models_info = []
        for m in models:
            if 'gemini' in m.name and 'generateContent' in m.supported_generation_methods:
                gemini_models_info.append({
                    "name": m.name,
                    "display_name": m.display_name,
                    "version": m.version,
                    "supported_generation_methods": m.supported_generation_methods,
                    "input_token_limit": m.input_token_limit,
                    "output_token_limit": m.output_token_limit
                })
        logger.info(f"Found {len(gemini_models_info)} Gemini models.")
        return {"available_gemini_models": gemini_models_info}
    except Exception as e:
        logger.error(f"Error listing Gemini models: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list Gemini models: {str(e)}")

@app.post("/ai/user-query")
async def ai_user_query(
    transcript: str = Form(...),
    user_prompt: str = Form(...),
    model: str = Form("claude-3-haiku-20240307"),
    max_tokens: int = Form(1000),
    user_plan: str = Form("free")
):
    logger.info(f"AI user query endpoint called. Model: {model}, Prompt: '{user_prompt}', User Plan: {user_plan}")

    if not is_paid_ai_user(user_plan):
        raise HTTPException(status_code=403, detail="AI Assistant features are only available for paid AI users (Three-Day, One-Week, Monthly Plan, Yearly Plan plans). Please upgrade your plan.")

    if not claude_client:
        raise HTTPException(status_code=503, detail=f"{TYPEMYWORDZ_AI_NAME} service is not initialized (API key missing or invalid).")

    try:
        if len(transcript) > 100000:
            raise HTTPException(status_code=400, detail="Transcript is too long. Please use a shorter transcript.")
        
        if len(user_prompt) > 1000:
            raise HTTPException(status_code=400, detail="User prompt is too long. Please use a shorter prompt.")

        full_prompt = f"{user_prompt}\n\nHere is the transcript:\n{transcript}"

        message = claude_client.messages.create(
            model=model,
            max_tokens=max_tokens,
            timeout=30.0,
            messages=[
                {"role": "user", "content": full_prompt}
            ]
        )
        ai_response = message.content[0].text
        logger.info(f"Successfully processed AI user query with {model}.")
        return {"ai_response": ai_response}

    except anthropic.APIError as e:
        error_message = "AI service error"
        error_details = str(e)
        
        if hasattr(e, 'body'):
            try:
                error_data = e.body if isinstance(e.body, dict) else {"error": str(e.body)}
                error_details = error_data
                logger.error(f"Anthropic API Error for user query: {error_data}")
            except:
                logger.error(f"Anthropic API Error for user query: {str(e)}")
        else:
            logger.error(f"Anthropic API Error for user query: {str(e)}")
            
        raise HTTPException(status_code=500, detail=f"{error_message}: {error_details}")
    
    except anthropic.APITimeoutError as e:
        logger.error(f"Anthropic API Timeout for user query: {e}")
        raise HTTPException(status_code=504, detail="AI service timeout. Please try again.")
    
    except Exception as e:
        logger.error(f"Unexpected error processing AI user query: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")

@app.post("/ai/admin-format")
async def ai_admin_format(
    transcript: str = Form(...),
    formatting_instructions: str = Form("Format the transcript for readability, correct grammar, and identify main sections with headings. Ensure a professional tone."),
    model: str = Form("claude-3-5-haiku-20241022"),
    max_tokens: int = Form(4000),
    user_plan: str = Form("free")
):
    logger.info(f"AI admin format endpoint (Anthropic) called. Model: {model}, Instructions: '{formatting_instructions}', User Plan: {user_plan}")

    if not is_paid_ai_user(user_plan):
        raise HTTPException(status_code=403, detail="AI Admin formatting features are only available for paid AI users (Three-Day, One-Week, Monthly Plan, Yearly Plan plans). Please upgrade your plan.")

    if not claude_client:
        raise HTTPException(status_code=503, detail=f"{TYPEMYWORDZ_AI_NAME} service is not initialized (API key missing or invalid).")

    try:
        if len(transcript) > 200000:
            raise HTTPException(status_code=400, detail="Transcript is too long. Please use a shorter transcript.")
        
        full_prompt = f"Please apply the following formatting and polishing instructions to the provided transcript:\n\nInstructions: {formatting_instructions}\n\nTranscript to format:\n{transcript}"

        message = claude_client.messages.create(
            model=model,
            max_tokens=max_tokens,
            timeout=60.0,
            messages=[
                {"role": "user", "content": full_prompt}
            ]
        )
        ai_response = message.content[0].text
        logger.info(f"Successfully processed AI admin format request with {model}.")
        return {"formatted_transcript": ai_response}

    except anthropic.APIError as e:
        error_message = "AI service error for admin formatting"
        error_details = str(e)
        
        if hasattr(e, 'body'):
            try:
                error_data = e.body if isinstance(e.body, dict) else {"error": str(e.body)}
                error_details = error_data
                logger.error(f"Anthropic API Error for admin format: {error_data}")
            except:
                logger.error(f"Anthropic API Error for admin format: {str(e)}")
        else:
            logger.error(f"Anthropic API Error for admin format: {str(e)}")
            
        raise HTTPException(status_code=500, detail=f"{error_message}: {error_details}")
    
    except anthropic.APITimeoutError as e:
        logger.error(f"Anthropic API Timeout for admin format: {e}")
        raise HTTPException(status_code=504, detail="AI service timeout. Please try again with a shorter transcript.")
    
    except Exception as e:
        logger.error(f"Unexpected error processing AI admin format request: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during admin formatting: {str(e)}")

@app.post("/ai/admin-format-gemini")
async def ai_admin_format_gemini(
    transcript: str = Form(...),
    formatting_instructions: str = Form("Correct all grammar, ensure a formal tone, break into paragraphs with subheadings for each major topic, and highlight action items in bold."),
    model: str = Form("models/gemini-pro-latest"),
    max_tokens: int = Form(4000),
    user_plan: str = Form("free")
):
    logger.info(f"AI admin format endpoint (Gemini) called. Model: {model}, Instructions: '{formatting_instructions}', User Plan: {user_plan}")

    if not is_paid_ai_user(user_plan):
        raise HTTPException(status_code=403, detail="AI Admin formatting features are only available for paid AI users (Three-Day, One-Week, Monthly Plan, Yearly Plan plans). Please upgrade your plan.")

    if not gemini_client:
        logger.error("Gemini client is not initialized in /ai/admin-format-gemini. Check GEMINI_API_KEY.")
        raise HTTPException(status_code=503, detail=f"Google Gemini service is not initialized (API key missing or invalid).")

    try:
        if len(transcript) > 200000:
            raise HTTPException(status_code=400, detail="Transcript is too long. Please use a shorter transcript.")
        
        full_prompt = f"Please apply the following formatting and polishing instructions to the provided transcript:\n\nInstructions: {formatting_instructions}\n\nTranscript to format:\n{transcript}"

        response = gemini_client.generate_content(
            full_prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=max_tokens,
                temperature=0.7,
                top_p=0.95,
                top_k=40
            )
        )
        # Handle safety filter or empty response case
        if not response.candidates or not response.candidates[0].content or not response.candidates[0].content.parts:
            logger.warning(f"Gemini response was filtered or empty. Finish reason: {response.candidates[0].finish_reason if response.candidates else 'unknown'}")
            return {"formatted_transcript": "The AI was unable to process this content due to content safety filters. Please try reformulating your request or use Claude instead."}
        
        gemini_response = response.text
        logger.info(f"Successfully processed AI admin format request with Gemini model: {model}.")
        return {"formatted_transcript": gemini_response}

    except Exception as e:
        logger.error(f"Unexpected error processing AI admin format request with Gemini: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during Gemini formatting: {str(e)}. Try using Claude instead.")

@app.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    language_code: Optional[str] = Form("en"),
    speaker_labels_enabled: bool = Form(False),
    user_plan: str = Form("free"),
    user_email: str = Form(""),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    logger.info(f"Main transcription endpoint called with file: {file.filename}, language: {language_code}, speaker_labels: {speaker_labels_enabled}, user_plan: {user_plan}, user_email: {user_email}")
    
    if not file.content_type.startswith(('audio/', 'video/')):
        logger.warning(f"Invalid file type: {file.content_type}")
        raise HTTPException(status_code=400, detail="Please upload an audio or video file")
    
    job_id = str(uuid.uuid4())
    logger.info(f"Created transcription job ID: {job_id}")
    
    try:
        logger.info(f"Saving uploaded file {file.filename} temporarily...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        file_size_mb = len(content) / (1024 * 1024)
        
        # Analyze audio to get duration
        audio_characteristics = await analyze_audio_characteristics(tmp_path)
        duration_seconds = audio_characteristics.get("duration_seconds", 0)
        duration_minutes = duration_seconds / 60.0
        
        jobs[job_id] = {
            "status": "processing",
            "filename": file.filename,
            "created_at": datetime.now().isoformat(),
            "file_size_mb": round(file_size_mb, 2),
            "content_type": file.content_type,
            "requested_language": language_code,
            "speaker_labels_enabled": speaker_labels_enabled,
            "user_plan": user_plan,
            "user_email": user_email,
            "duration_minutes": duration_minutes,
            "duration_seconds": duration_seconds
        }
        
        cancellation_flags[job_id] = False
        logger.info(f"Job {job_id} initialized with status 'processing'")
        
    except Exception as e:
        logger.error(f"ERROR processing file for job {job_id}: {str(e)}")
        if job_id in jobs:
            del jobs[job_id]
        if job_id in cancellation_flags:
            del cancellation_flags[job_id]
        raise HTTPException(status_code=500, detail="Failed to process audio file")

    background_tasks.add_task(
        process_transcription_job, 
        job_id, 
        tmp_path, 
        file.filename, 
        language_code, 
        speaker_labels_enabled, 
        user_plan, 
        duration_minutes,
        user_email
    )
    
    logger.info(f"Returning immediate response for job ID: {job_id}")
    return {
        "job_id": job_id, 
        "status": jobs[job_id]["status"],
        "filename": file.filename,
        "file_size_mb": jobs[job_id]["file_size_mb"],
        "duration_minutes": duration_minutes,
        "created_at": jobs[job_id]["created_at"],
        "logic_used": f"UserPlan:{user_plan}, Email:{user_email}, Admin:{is_admin_user(user_email)}"
    }

@app.post("/generate-formatted-word")
async def generate_formatted_word(request: FormattedWordDownloadRequest):
    logger.info(f"Generating formatted Word document for {request.filename}")
    try:
        document = Document()
        lines = request.transcription_html.split('\n')
        
        speaker_tag_pattern = re.compile(r'<strong>(Speaker \d+:)</strong>(.*)')
        
        for line in lines:
            if line.strip():
                p = document.add_paragraph()
                
                match = speaker_tag_pattern.match(line)
                if match:
                    speaker_label_text = match.group(1)
                    remaining_text = match.group(2).strip()

                    run = p.add_run(speaker_label_text)
                    run.bold = True
                    
                    if remaining_text:
                        p.add_run(" " + remaining_text)
                else:
                    clean_line = re.sub(r'<[^>]*>', '', line).strip()
                    if clean_line:
                        p.add_run(clean_line)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={request.filename}"}
        )

    except Exception as e:
        logger.error(f"Error generating formatted Word document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate formatted Word document: {str(e)}")

@app.get("/status/{job_id}")
async def get_job_status(job_id: str):
    logger.info(f"Status check for job ID: {job_id}")
    if job_id not in jobs:
        logger.warning(f"Job ID {job_id} not found")
        raise HTTPException(status_code=404, detail="Job not found")
    
    job_data = jobs[job_id]
    
    if job_data["status"] == "cancelled" or cancellation_flags.get(job_id, False):
        logger.info(f"Job {job_id} was cancelled, returning cancelled status")
        job_data["status"] = "cancelled"
    
    return job_data

@app.post("/cancel/{job_id}")
async def cancel_job(job_id: str):
    logger.info(f"Cancel request received for job ID: {job_id}")
    
    if job_id not in jobs:
        logger.warning(f"Cancel request: Job ID {job_id} not found")
        raise HTTPException(status_code=404, detail="Job not found")
    
    job_data = jobs[job_id]
    
    try:
        cancellation_flags[job_id] = True
        logger.info(f"Cancellation flag set for job {job_id}")
        
        if job_id in active_background_tasks:
            task = active_background_tasks[job_id]
            if not task.done():
                logger.info(f"Cancelling active background task for job {job_id}")
                task.cancel()
                try:
                    await asyncio.wait_for(task, timeout=2.0)
                except (asyncio.TimeoutError, asyncio.CancelledError):
                    logger.info(f"Background task for job {job_id} cancelled (timeout/cancelled)")
            else:
                logger.info(f"Background task for job {job_id} was already completed")
        
        job_data.update({
            "status": "cancelled",
            "cancelled_at": datetime.now().isoformat(),
            "error": "Job was cancelled by user"
        })
        
        logger.info(f"Job {job_id} successfully cancelled")
        return {
            "message": "Job cancelled successfully", 
            "job_id": job_id,
            "cancelled_at": job_data["cancelled_at"],
            "previous_status": job_data.get("previous_status", "processing")
        }
        
    except Exception as e:
        logger.error(f"Error cancelling job {job_id}: {str(e)}")
        job_data.update({
            "status": "cancelled",
            "cancelled_at": datetime.now().isoformat(),
            "error": f"Job cancelled with errors: {str(e)}"
        })
        raise HTTPException(status_code=500, detail=f"Job cancelled but with errors: {str(e)}")

@app.post("/compress-download")
async def compress_download(file: UploadFile = File(...), quality: str = "high"):
    """Endpoint to compress audio files for download"""
    logger.info(f"Compress download endpoint called with file: {file.filename}, quality: {quality}")
    
    if not file.content_type.startswith(('audio/', 'video/')):
        logger.warning(f"Invalid file type: {file.content_type}")
        raise HTTPException(status_code=400, detail="Please upload an audio or video file")
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            input_path = tmp.name
        
        output_path = compress_audio_for_download(input_path, quality=quality)
        
        with open(output_path, 'rb') as f:
            compressed_content = f.read()
        
        os.unlink(input_path)
        os.unlink(output_path)
        
        from fastapi.responses import Response as FastAPIResponse
        return FastAPIResponse(
            content=compressed_content,
            media_type="audio/mp3",
            headers={"Content-Disposition": f"attachment; filename=compressed_{file.filename}.mp3"}
        )
        
    except Exception as e:
        logger.error(f"Error compressing file for download: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to compress audio file: {str(e)}")

@app.get("/jobs")
async def list_jobs():
    logger.info("Jobs list endpoint called")
    
    job_summary = {}
    for job_id, job_data in jobs.items():
        job_summary[job_id] = {
            "status": job_data["status"],
            "filename": job_data.get("filename", "unknown"),
            "created_at": datetime.fromisoformat(job_data["created_at"]).strftime('%Y-%m-%d %H:%M:%S'),
            "file_size_mb": job_data.get("file_size_mb", 0),
            "duration_minutes": job_data.get("duration_minutes", 0),
            "user_plan": job_data.get("user_plan", "unknown"),
            "user_email": job_data.get("user_email", ""),
            "is_admin": is_admin_user(job_data.get("user_email", "")),
            "primary_service": job_data.get("tier_1_service"),
            "service_used": (job_data.get("tier_1_used") or job_data.get("tier_2_used") or job_data.get("tier_3_used")),
            "model_used": job_data.get("model_used", "N/A"),
            "has_background_task": job_id in active_background_tasks,
            "is_cancellation_flagged": cancellation_flags.get(job_id, False),
            "word_count": job_data.get("word_count"),
            "duration_seconds": job_data.get("duration_seconds"),
            "requested_language": job_data.get("requested_language", "en"),
            "speaker_labels_enabled": job_data.get("speaker_labels_enabled", False),
            "selection_reason": job_data.get("selection_reason", "unknown")
        }
    
    return {
        "total_jobs": len(jobs),
        "active_background_tasks": len(active_background_tasks),
        "cancellation_flags": len(cancellation_flags),
        "jobs": job_summary,
        "admin_emails": ADMIN_EMAILS,
        "assemblyai_tester_email": ASSEMBLYAI_TESTER_EMAIL,
        "gemini_access": "NOW AVAILABLE FOR ALL PAID AI USERS (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)",
        "system_stats": {
            "jobs_by_status": {
                status: len([j for j in jobs.values() if j["status"] == status])
                for status in ["processing", "completed", "failed", "cancelled"]
            }
        }
    }

@app.get("/health")
async def health_check():
    logger.info("Health check endpoint called")
    
    try:
        import psutil
        memory_info = psutil.virtual_memory()
        cpu_percent = psutil.cpu_percent(interval=1)
        
        health_data = {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "system": {
                "memory_percent": memory_info.percent,
                "cpu_percent": cpu_percent,
                "available_ram_gb": round(memory_info.available / (1024**3), 2)
            },
            "application": {
                "total_jobs": len(jobs),
                "active_background_tasks": len(active_background_tasks),
                "cancellation_flags": len(cancellation_flags),
                "jobs_by_status": {
                    status: len([j for j in jobs.values() if j["status"] == status])
                    for status in ["processing", "completed", "failed", "cancelled"]
                }
            },
            "integrations": {
                "assemblyai_configured": bool(ASSEMBLYAI_API_KEY),
                "anthropic_configured": bool(ANTHROPIC_API_KEY),
                "openai_configured": bool(OPENAI_API_KEY),
                "openai_whisper_service_configured": bool(OPENAI_WHISPER_SERVICE_RAILWAY_URL),
                "google_gemini_configured": bool(GEMINI_API_KEY),
                "deepgram_service_configured": bool(DEEPGRAM_SERVICE_RAILWAY_URL)
            },
            "transcription_logic": {
                "free_user_transcription": f"Primary={TYPEMYWORDZ1_NAME} → Fallback1={DEEPGRAM_NAME} → Fallback2=None",
                "three_day_plan_transcription": f"Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}",
                "one_week_plan_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
                "monthly_plan_transcription": f"Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}",
                "yearly_plan_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
                "admin_transcription": f"Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}",
                "speaker_labels_transcription": f"Always use {TYPEMYWORDZ1_NAME} first → Fallback1={DEEPGRAM_NAME} → Fallback2=None",
                "assemblyai_tester_transcription": f"Always use {TYPEMYWORDZ1_NAME} (no fallback for {ASSEMBLYAI_TESTER_EMAIL})",
                "free_users_assemblyai_model": f"{TYPEMYWORDZ1_NAME} nano model",
                "paid_users_assemblyai_model": f"{TYPEMYWORDZ1_NAME} best model",
                "ai_features_access": "Only for Three-Day, One-Week, Monthly Plan, and Yearly Plan plans",
                "gemini_access": "NOW AVAILABLE FOR ALL PAID AI USERS (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)",
                "assemblyai": f"TypeMyworDz1 (AssemblyAI)",
                "openai_whisper": f"TypeMyworDz2 (OpenAI Whisper-1)",
                "deepgram": f"Deepgram",
                "anthropic_ai": f"TypeMyworDz AI (Anthropic Claude)",
                "google_gemini_ai": "Google Gemini - Available for ALL paid AI users",
                "admin_emails": ADMIN_EMAILS,
                "assemblyai_tester_email": ASSEMBLYAI_TESTER_EMAIL
            }
        }
        
        return health_data
        
    except Exception as e:
        logger.error(f"Health check error: {e}")
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

logger.info("=== FASTAPI APPLICATION SETUP COMPLETE ===")

logger.info("Performing final system validation...")
logger.info(f"TypeMyworDz1 API Key configured: {bool(ASSEMBLYAI_API_KEY)}")
logger.info(f"TypeMyworDz2 Service URL configured: {bool(OPENAI_WHISPER_SERVICE_RAILWAY_URL)}")
logger.info(f"Deepgram Service URL configured: {bool(DEEPGRAM_SERVICE_RAILWAY_URL)}")
logger.info(f"TypeMyworDz AI API Key configured: {bool(ANTHROPIC_API_KEY)}")
logger.info(f"OpenAI GPT API Key configured: {bool(OPENAI_API_KEY)}")
logger.info(f"Google Gemini API Key configured: {bool(GEMINI_API_KEY)}")
logger.info(f"Paystack Secret Key configured: {bool(PAYSTACK_SECRET_KEY)}")
logger.info(f"Firebase Admin SDK configured: {bool(FIREBASE_ADMIN_SDK_CONFIG_BASE64) and bool(db)}")
logger.info(f"Admin emails configured: {ADMIN_EMAILS}")
logger.info(f"AssemblyAI Tester email configured: {ASSEMBLYAI_TESTER_EMAIL}")
logger.info(f"UPDATED: Google Gemini now available for ALL PAID AI USERS (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)")
logger.info(f"Job tracking systems initialized:")
logger.info(f"  - Main jobs dictionary: {len(jobs)} jobs")
logger.info(f"  - Active background tasks: {len(active_background_tasks)} tasks")
logger.info(f"  - Cancellation flags: {len(cancellation_flags)} flags")
logger.info("Available API endpoints:")
logger.info("  POST /transcribe - Main transcription endpoint with smart service selection")
logger.info("  POST /ai/user-query - Process user-driven AI queries (summarize, Q&A, bullet points) with Claude")
logger.info("  POST /ai/user-query-gemini - Process user-driven AI queries with Gemini (NOW FOR ALL PAID USERS)")
logger.info("  POST /ai/admin-format - Process admin-driven AI formatting requests (Anthropic)")
logger.info("  POST /ai/admin-format-gemini - Process admin-driven AI formatting requests (Google Gemini - NOW FOR ALL PAID USERS)")
logger.info("  POST /api/initialize-paystack-payment - Initialize Paystack payment")
logger.info("  POST /api/verify-payment - Verify Paystack payment")
logger.info("  POST /api/paystack-webhook - Handle Paystack webhooks")
logger.info("  GET /api/paystack-status - Get integration status")
logger.info("  GET /api/list-gemini-models - List available Gemini models")
logger.info("  GET /status/{job_id} - Check job status")
logger.info("  POST /cancel/{job_id} - Cancel transcription job")
logger.info("  POST /compress-download - Compress audio for download")
logger.info("  POST /generate-formatted-word - Generate formatted Word document with speaker labels")
logger.info("  GET /jobs - List all jobs")
logger.info("  GET /health - System health check")
logger.info("  DELETE /cleanup - Clean up old jobs")
logger.info("  GET / - Root endpoint with service info")

if __name__ == "__main__":
    logger.info("Starting Uvicorn server directly...")
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    host = os.environ.get("HOST", "0.0.0.0")
    
    logger.info(f"Starting enhanced transcription service on {host}:{port}")
    logger.info("🚀 NEW ENHANCED FEATURES:")
    logger.info(f"  ✅ Smart service selection with updated three-tier logic")
    logger.info(f"  ✅ Three-tier automatic fallback system")
    logger.info(f"  ✅ Admin email-based service prioritization")
    logger.info(f"  ✅ Dedicated AssemblyAI tester logic")
    logger.info(f"  ✅ Speaker diarization for AssemblyAI and Deepgram")
    logger.info(f"  ✅ Dynamic TypeMyworDz1 model selection (nano for free, best for paid)")
    logger.info("  ✅ Unified transcription processing pipeline")
    logger.info("  ✅ Enhanced error handling and service resilience")
    logger.info("  ✅ Paystack payment integration")
    logger.info("  ✅ Multi-language support")
    logger.info("  ✅ Formatted Word document generation")
    logger.info(f"  ✅ User-driven AI features (summarization, Q&A, and bullet points) via TypeMyworDz AI (Anthropic)")
    logger.info(f"  ✅ Admin-driven AI formatting via TypeMyworDz AI (Anthropic) and Google Gemini")
    logger.info(f"  ✅ Google Gemini integration for AI queries - NOW AVAILABLE FOR ALL PAID AI USERS")
    logger.info(f"  ✅ AI Assistant features restricted to paid users (Three-Day, One-Week, Monthly Plan, Yearly Plan plans)")
    logger.info("  🆕 UPDATED: Google Gemini now accessible to ALL paid AI users, not just admins")
    
    logger.info("🔧 NEW TRANSCRIPTION LOGIC:")
    logger.info(f"  - Free users: Primary={TYPEMYWORDZ1_NAME} → Fallback1={DEEPGRAM_NAME} → Fallback2=None")
    logger.info(f"  - Three-Day Plan: Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}")
    logger.info(f"  - One-Week Plan: Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}")
    logger.info(f"  - Monthly Plan: Primary={DEEPGRAM_NAME} → Fallback1={TYPEMYWORDZ2_NAME} → Fallback2={TYPEMYWORDZ1_NAME}")
    logger.info(f"  - Yearly Plan & Admins ({', '.join(ADMIN_EMAILS)}): Primary={TYPEMYWORDZ2_NAME} → Fallback1={TYPEMYWORDZ1_NAME} → Fallback2={DEEPGRAM_NAME}")
    logger.info(f"  - Speaker Labels requested: Always use {TYPEMYWORDZ1_NAME} first → Fallback1={DEEPGRAM_NAME} → Fallback2=None")
    logger.info(f"  - Dedicated AssemblyAI Tester ({ASSEMBLYAI_TESTER_EMAIL}): Primary={TYPEMYWORDZ1_NAME} (no fallback)")
    logger.info(f"  - Free users: {TYPEMYWORDZ1_NAME} nano model")
    logger.info(f"  - Paid users: {TYPEMYWORDZ1_NAME} best model")
    logger.info(f"  - {TYPEMYWORDZ1_NAME}: AssemblyAI")
    logger.info(f"  - {TYPEMYWORDZ2_NAME}: OpenAI Whisper-1 (typically does NOT support speaker labels)")
    logger.info(f"  - {DEEPGRAM_NAME}: Deepgram")
    logger.info(f"  - {TYPEMYWORDZ_AI_NAME} (Anthropic Claude 3 Haiku / 3.5 Haiku) for user AI text processing")
    logger.info(f"  - Google Gemini for AI text processing - NOW AVAILABLE FOR ALL PAID AI USERS")
    
    try:
        uvicorn.run(
            app, 
            host=host, 
            port=port,
            log_level="info",
            access_log=True,
            reload=False,
            workers=1
        )
    except Exception as e:
        logger.error(f"Failed to start server: {e}")
        sys.exit(1)
else:
    logger.info("Application loaded as module")
    logger.info(f"Ready to handle requests with {TYPEMYWORDZ1_NAME} + {TYPEMYWORDZ2_NAME} + {DEEPGRAM_NAME} + {TYPEMYWORDZ_AI_NAME} (Anthropic) + Google Gemini integration")
